"""
Génère les fichiers .docx Jinja2 utilisés comme templates par docxtpl.
Lancer une fois depuis le répertoire backend/ :
    python scripts/create_templates.py
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)

LEXAFRICA_GRAY = RGBColor(0x55, 0x55, 0x55)
LEXAFRICA_BLUE = RGBColor(0x1A, 0x56, 0xDB)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _add_watermark_header(doc, text="APERÇU - LexAfrica"):
    """Ajoute un en-tête discret servant de filigrane texte."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run.font.italic = True


def _add_logo_header(doc, title):
    """En-tête avec nom de la plateforme et titre du document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LexAfrica")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = LEXAFRICA_BLUE
    doc.add_paragraph()

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_article(doc, number, title, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"ARTICLE {number} – {title}")
    run.bold = True
    run.font.size = Pt(11)
    body = doc.add_paragraph(content)
    body.paragraph_format.left_indent = Pt(12)
    doc.add_paragraph()


def _add_footer(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(
        "─────────────────────────────────────────────\n"
        "Document généré par LexAfrica · www.lexafrica.com\n"
        "Modèles juridiques validés par nos partenaires avocats OHADA\n"
        "Ce document est un aperçu. Effectuez le paiement pour télécharger la version définitive."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = LEXAFRICA_GRAY
    run.font.italic = True


def _add_signature_block(doc, left_title, left_name, right_title, right_name):
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    cells = table.rows

    def cell_text(row, col, text, bold=False):
        c = table.cell(row, col)
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(text)
        run.bold = bold
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_text(0, 0, left_title, bold=True)
    cell_text(0, 1, right_title, bold=True)
    cell_text(1, 0, left_name)
    cell_text(1, 1, right_name)
    cell_text(2, 0, "")
    cell_text(2, 1, "")
    cell_text(3, 0, "Signature et cachet")
    cell_text(3, 1, "Signature (précédée de\n« Lu et approuvé »)")
    doc.add_paragraph()


# ── Templates ──────────────────────────────────────────────────────────────────

def create_cdi_template(country_code: str):
    doc = Document()
    _set_margins(doc)
    _add_watermark_header(doc)
    _add_logo_header(doc, "CONTRAT DE TRAVAIL À DURÉE INDÉTERMINÉE\n(CDI – {{ country_full_name }})")

    # Parties
    doc.add_heading("ENTRE LES SOUSSIGNÉS :", level=2)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("L'EMPLOYEUR\n").bold = True
    p.add_run(
        "{{ employer_name }}, dont le siège social est situé à {{ employer_address }},"
        "{% if employer_rccm %} immatriculée au RCCM sous le numéro {{ employer_rccm }},{% endif %}\n"
        "représentée par {{ employer_representative|default('son représentant légal') }},\n"
        "ci-après dénommée « L'Employeur »,"
    )

    doc.add_paragraph()
    doc.add_paragraph("ET").alignment

    p = doc.add_paragraph()
    p.add_run("L'EMPLOYÉ(E)\n").bold = True
    p.add_run(
        "Monsieur/Madame {{ employee_name }}, né(e) le {{ employee_birth_date }}"
        "{% if employee_birth_place %} à {{ employee_birth_place }}{% endif %}, "
        "de nationalité {{ employee_nationality|default('ivoirienne') }},\n"
        "demeurant à {{ employee_address }},\n"
        "ci-après dénommé(e) « L'Employé(e) »,"
    )

    doc.add_paragraph()
    bold_p = doc.add_paragraph()
    bold_p.add_run("IL A ÉTÉ CONVENU ET ARRÊTÉ CE QUI SUIT :").bold = True
    doc.add_paragraph()

    _add_article(doc, 1, "ENGAGEMENT",
        "L'Employeur engage Monsieur/Madame {{ employee_name }} en qualité de "
        "{{ job_title }} (catégorie : {{ job_category }}) à compter du {{ start_date }}."
    )
    _add_article(doc, 2, "PÉRIODE D'ESSAI",
        "Le présent contrat est soumis à une période d'essai de {{ trial_period_months }} mois, "
        "renouvelable une fois dans la limite légale de {{ trial_period_legal_max }} mois, "
        "conformément au {{ country_code_travail }}.\n"
        "Durant cette période, chacune des parties peut mettre fin au contrat sans préavis "
        "durant le premier mois, puis avec un préavis de 8 jours."
    )
    _add_article(doc, 3, "RÉMUNÉRATION",
        "En contrepartie de ses services, l'Employé(e) percevra une rémunération mensuelle brute "
        "de {{ salary }} {{ currency }}, versée à terme échu.\n"
        "Ce salaire est supérieur ou égal au SMIG en vigueur en {{ country_full_name }} "
        "({{ smig_label }})."
    )
    _add_article(doc, 4, "LIEU DE TRAVAIL",
        "Le lieu habituel de travail est fixé à {{ work_location }}. "
        "Toute modification substantielle du lieu de travail fera l'objet d'un avenant "
        "signé par les deux parties."
    )
    _add_article(doc, 5, "DURÉE ET ORGANISATION DU TRAVAIL",
        "La durée hebdomadaire de travail est de {{ weekly_hours }} heures, conformément "
        "au {{ country_code_travail }}.\n"
        "Les heures supplémentaires éventuelles seront rémunérées aux taux légaux en vigueur : "
        "{{ overtime_rate_1 }} pour les 8 premières heures au-delà de la durée légale, "
        "{{ overtime_rate_2 }} au-delà."
    )
    _add_article(doc, 6, "CONGÉS PAYÉS",
        "L'Employé(e) bénéficiera de {{ paid_leave_days_per_year }} jours ouvrables "
        "de congés payés par an ({{ paid_leave_days_per_month }} jours par mois travaillé), "
        "conformément au {{ country_code_travail }}."
    )
    _add_article(doc, 7, "OBLIGATIONS DE L'EMPLOYÉ(E)",
        "L'Employé(e) s'engage à :\n"
        "- Exercer ses fonctions avec diligence, loyauté et professionnalisme ;\n"
        "- Respecter le règlement intérieur de l'Employeur ;\n"
        "- Maintenir la confidentialité sur les informations de l'entreprise ;\n"
        "- Signaler immédiatement tout conflit d'intérêts potentiel."
    )
    _add_article(doc, 8, "PROTECTION SOCIALE",
        "L'Employeur procédera à l'affiliation de l'Employé(e) à la "
        "{{ social_protection_body }} et s'acquittera des cotisations patronales et salariales "
        "conformément à la {{ social_protection_law }}."
    )
    _add_article(doc, 9, "PRÉAVIS ET RUPTURE DU CONTRAT",
        "Le présent contrat pourra être rompu par l'une ou l'autre des parties sous réserve du "
        "respect d'un préavis de {{ notice_period_months }} mois, sauf faute grave ou lourde "
        "dûment constatée, ou accord mutuel des parties."
    )
    _add_article(doc, 10, "LOI APPLICABLE ET JURIDICTION",
        "Le présent contrat est régi par le {{ country_code_travail }} ainsi que par "
        "l'Acte Uniforme OHADA relatif au droit du travail. "
        "Tout litige relatif à son exécution ou à sa résiliation sera soumis au "
        "{{ jurisdiction }}."
    )

    doc.add_paragraph(
        "Fait en deux (2) exemplaires originaux, "
        "à {{ signature_place|default('Abidjan') }}, le {{ signature_date }}."
    )

    _add_signature_block(
        doc,
        "Pour l'Employeur", "{{ employer_name }}\n{{ employer_representative|default('') }}",
        "L'Employé(e)", "{{ employee_name }}"
    )
    _add_footer(doc)

    filename = f"cdi_{country_code.lower()}.docx"
    doc.save(os.path.join(TEMPLATES_DIR, filename))
    print(f"  ✓ {filename}")


def create_cdd_template(country_code: str):
    doc = Document()
    _set_margins(doc)
    _add_watermark_header(doc)
    _add_logo_header(doc, "CONTRAT DE TRAVAIL À DURÉE DÉTERMINÉE\n(CDD – {{ country_full_name }})")

    doc.add_heading("ENTRE LES SOUSSIGNÉS :", level=2)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("L'EMPLOYEUR\n").bold = True
    p.add_run(
        "{{ employer_name }}, dont le siège social est situé à {{ employer_address }},"
        "{% if employer_rccm %} immatriculée au RCCM sous le numéro {{ employer_rccm }},{% endif %}\n"
        "représentée par {{ employer_representative|default('son représentant légal') }},\n"
        "ci-après dénommée « L'Employeur »,"
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("L'EMPLOYÉ(E)\n").bold = True
    p.add_run(
        "Monsieur/Madame {{ employee_name }}, né(e) le {{ employee_birth_date }}"
        "{% if employee_birth_place %} à {{ employee_birth_place }}{% endif %}, "
        "de nationalité {{ employee_nationality|default('ivoirienne') }},\n"
        "demeurant à {{ employee_address }},\n"
        "ci-après dénommé(e) « L'Employé(e) »,"
    )

    doc.add_paragraph()
    bp = doc.add_paragraph()
    bp.add_run("IL A ÉTÉ CONVENU ET ARRÊTÉ CE QUI SUIT :").bold = True
    doc.add_paragraph()

    _add_article(doc, 1, "NATURE ET OBJET DU CONTRAT",
        "Le présent contrat est conclu pour une durée déterminée pour le motif suivant : "
        "{{ contract_purpose }}.\n"
        "Conformément au {{ country_code_travail }}, ce contrat ne peut être utilisé "
        "pour pourvoir durablement un emploi lié à l'activité normale et permanente de l'entreprise."
    )
    _add_article(doc, 2, "DURÉE",
        "Le présent contrat prend effet le {{ start_date }} et prend fin le {{ end_date }}, "
        "soit une durée totale de {{ contract_duration|default('à calculer') }}."
    )
    _add_article(doc, 3, "PÉRIODE D'ESSAI",
        "Le présent contrat est soumis à une période d'essai de {{ trial_period_months }} mois "
        "conformément au {{ country_code_travail }}."
    )
    _add_article(doc, 4, "POSTE ET RÉMUNÉRATION",
        "L'Employé(e) est engagé(e) en qualité de {{ job_title }} (catégorie : {{ job_category }}) "
        "au lieu de travail de {{ work_location }}, pour une rémunération mensuelle brute de "
        "{{ salary }} {{ currency }}."
    )
    _add_article(doc, 5, "DURÉE DU TRAVAIL",
        "La durée hebdomadaire de travail est de {{ weekly_hours }} heures, "
        "conformément au {{ country_code_travail }}."
    )
    _add_article(doc, 6, "RENOUVELLEMENT",
        "Le présent CDD peut être renouvelé une fois, dans les conditions et limites "
        "prévues par le {{ country_code_travail }}, sous réserve d'un accord écrit "
        "avant le terme initial."
    )
    _add_article(doc, 7, "PROTECTION SOCIALE",
        "L'Employeur procédera à l'affiliation de l'Employé(e) à la "
        "{{ social_protection_body }} pour la durée du contrat."
    )
    _add_article(doc, 8, "LOI APPLICABLE",
        "Le présent contrat est régi par le {{ country_code_travail }}. "
        "Tout litige sera soumis au {{ jurisdiction }}."
    )

    doc.add_paragraph(
        "Fait en deux (2) exemplaires originaux, "
        "à {{ signature_place|default('Abidjan') }}, le {{ signature_date }}."
    )

    _add_signature_block(
        doc,
        "Pour l'Employeur", "{{ employer_name }}\n{{ employer_representative|default('') }}",
        "L'Employé(e)", "{{ employee_name }}"
    )
    _add_footer(doc)

    filename = f"cdd_{country_code.lower()}.docx"
    doc.save(os.path.join(TEMPLATES_DIR, filename))
    print(f"  ✓ {filename}")


def create_nda_template():
    doc = Document()
    _set_margins(doc)
    _add_watermark_header(doc)
    _add_logo_header(doc, "ACCORD DE CONFIDENTIALITÉ\n(Non-Disclosure Agreement – NDA)")

    doc.add_heading("ENTRE LES SOUSSIGNÉS :", level=2)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("LA PARTIE DIVULGATRICE\n").bold = True
    p.add_run(
        "{{ party1_name }}, {{ party1_type|default('société') }} dont le siège est à "
        "{{ party1_address }},\n"
        "représentée par {{ party1_representative }},\n"
        "ci-après dénommée « La Partie Divulgatrice »,"
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("LA PARTIE RÉCEPTRICE\n").bold = True
    p.add_run(
        "{{ party2_name }}, {{ party2_type|default('société') }} dont le siège est à "
        "{{ party2_address }},\n"
        "représentée par {{ party2_representative }},\n"
        "ci-après dénommée « La Partie Réceptrice »,"
    )

    doc.add_paragraph()
    bp = doc.add_paragraph()
    bp.add_run("ONT CONVENU CE QUI SUIT :").bold = True
    doc.add_paragraph()

    _add_article(doc, 1, "OBJET",
        "Dans le cadre de {{ agreement_subject }}, la Partie Divulgatrice est susceptible "
        "de communiquer à la Partie Réceptrice des informations confidentielles.\n"
        "Le présent accord a pour objet de définir les conditions dans lesquelles ces "
        "informations seront protégées."
    )
    _add_article(doc, 2, "DÉFINITION DES INFORMATIONS CONFIDENTIELLES",
        "Sont considérées comme confidentielles toutes les informations, données, documents, "
        "fichiers, procédés, méthodes, et savoir-faire communiqués par la Partie Divulgatrice, "
        "sous quelque forme que ce soit (écrite, orale, électronique, visuelle), "
        "qu'ils soient ou non identifiés comme tels.\n"
        "Sont exclus : les informations déjà publiques, celles obtenues de tiers autorisés, "
        "ou celles dont la Partie Réceptrice démontre qu'elle les connaissait avant la divulgation."
    )
    _add_article(doc, 3, "OBLIGATIONS DE LA PARTIE RÉCEPTRICE",
        "La Partie Réceptrice s'engage à :\n"
        "- Maintenir les informations confidentielles strictement secrètes ;\n"
        "- Ne les utiliser qu'aux fins de {{ agreement_subject }} ;\n"
        "- Ne les divulguer qu'aux membres de son personnel en ayant strictement besoin ;\n"
        "- Mettre en place des mesures de protection équivalentes à celles qu'elle applique "
        "à ses propres informations confidentielles."
    )
    _add_article(doc, 4, "DURÉE",
        "Le présent accord entre en vigueur à la date de signature et reste en vigueur "
        "pendant {{ duration_years }} ans.\n"
        "Les obligations de confidentialité subsistent pendant {{ duration_years }} ans "
        "après l'expiration ou la résiliation du présent accord."
    )
    _add_article(doc, 5, "PROPRIÉTÉ INTELLECTUELLE",
        "La communication d'informations confidentielles ne confère à la Partie Réceptrice "
        "aucun droit de propriété intellectuelle sur lesdites informations. "
        "Celles-ci restent la propriété exclusive de la Partie Divulgatrice."
    )
    _add_article(doc, 6, "LOI APPLICABLE ET JURIDICTION",
        "Le présent accord est régi par le {{ governing_law_label }}. "
        "Tout litige sera soumis au {{ jurisdiction }}, après tentative de règlement amiable."
    )

    doc.add_paragraph(
        "Fait en deux (2) exemplaires originaux, "
        "à {{ signature_place|default('Abidjan') }}, le {{ signature_date }}."
    )

    _add_signature_block(
        doc,
        "La Partie Divulgatrice", "{{ party1_name }}\n{{ party1_representative }}",
        "La Partie Réceptrice", "{{ party2_name }}\n{{ party2_representative }}"
    )
    _add_footer(doc)

    doc.save(os.path.join(TEMPLATES_DIR, "nda.docx"))
    print("  ✓ nda.docx")


def create_prestation_template():
    doc = Document()
    _set_margins(doc)
    _add_watermark_header(doc)
    _add_logo_header(doc, "CONTRAT DE PRESTATION DE SERVICES\n({{ country_full_name }})")

    doc.add_heading("ENTRE LES SOUSSIGNÉS :", level=2)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("LE CLIENT\n").bold = True
    p.add_run(
        "{{ client_name }}, dont le siège est à {{ client_address }},\n"
        "représenté par {{ client_representative }},\n"
        "ci-après dénommé « Le Client »,"
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("LE PRESTATAIRE\n").bold = True
    p.add_run(
        "{{ consultant_name }}, {{ consultant_type|default('') }} "
        "dont l'adresse est {{ consultant_address }},\n"
        "{% if consultant_expertise %}Expert en {{ consultant_expertise }},\n{% endif %}"
        "ci-après dénommé « Le Prestataire »,"
    )

    doc.add_paragraph()
    bp = doc.add_paragraph()
    bp.add_run("IL A ÉTÉ CONVENU CE QUI SUIT :").bold = True
    doc.add_paragraph()

    _add_article(doc, 1, "OBJET DE LA MISSION",
        "Le Client confie au Prestataire la mission suivante :\n"
        "{{ mission_description }}\n\n"
        "Livrables attendus : {{ deliverables }}"
    )
    _add_article(doc, 2, "DURÉE DE LA MISSION",
        "La mission commence le {{ mission_start_date }} et se termine le {{ mission_end_date }}. "
        "Toute prolongation devra faire l'objet d'un avenant écrit signé par les deux parties."
    )
    _add_article(doc, 3, "RÉMUNÉRATION ET MODALITÉS DE PAIEMENT",
        "En contrepartie de la mission, le Client versera au Prestataire :\n"
        "- Montant total : {{ total_amount }} {{ currency }} hors taxes\n"
        "- Modalités : {{ payment_terms }}\n\n"
        "{{ tva_note }}"
    )
    _add_article(doc, 4, "INDÉPENDANCE DU PRESTATAIRE",
        "Le Prestataire exerce sa mission en toute indépendance. "
        "Le présent contrat ne crée aucun lien de subordination ni de contrat de travail "
        "entre les parties. Le Prestataire demeure libre d'organiser son temps et ses méthodes "
        "de travail, sous réserve du respect des délais et des livrables convenus."
    )
    _add_article(doc, 5, "CONFIDENTIALITÉ",
        "Le Prestataire s'engage à maintenir strictement confidentielle toute information "
        "obtenue dans le cadre de cette mission, pendant et après l'exécution du contrat."
    )
    _add_article(doc, 6, "PROPRIÉTÉ DES LIVRABLES",
        "Sauf stipulation contraire, les livrables produits dans le cadre de cette mission "
        "deviennent la propriété exclusive du Client après paiement intégral de la rémunération."
    )
    _add_article(doc, 7, "LOI APPLICABLE",
        "Le présent contrat est régi par le {{ governing_law_label }}. "
        "Tout litige sera soumis au {{ jurisdiction }}."
    )

    doc.add_paragraph(
        "Fait en deux (2) exemplaires originaux, "
        "à {{ signature_place|default('Abidjan') }}, le {{ signature_date }}."
    )

    _add_signature_block(
        doc,
        "Le Client", "{{ client_name }}\n{{ client_representative }}",
        "Le Prestataire", "{{ consultant_name }}"
    )
    _add_footer(doc)

    doc.save(os.path.join(TEMPLATES_DIR, "prestation_services.docx"))
    print("  ✓ prestation_services.docx")


def create_pacte_associes_template():
    doc = Document()
    _set_margins(doc)
    _add_watermark_header(doc)
    _add_logo_header(doc, "PACTE D'ASSOCIÉS SIMPLIFIÉ\n({{ company_type }} – OHADA · {{ country_full_name }})")

    p = doc.add_paragraph("Relatif à la société : ")
    run = p.add_run("{{ company_name }}")
    run.bold = True

    doc.add_paragraph(
        "Siège social : {{ company_address }}\n"
        "Capital social : {{ share_capital }} {{ currency }}\n"
        "Immatriculée au {{ company_registry }}"
    )
    doc.add_paragraph()
    bp = doc.add_paragraph()
    bp.add_run("ENTRE LES ASSOCIÉS SOUSSIGNÉS :").bold = True

    doc.add_paragraph(
        "{% for a in associates %}"
        "- {{ a.name }}, détenant {{ a.shares_percentage }}% du capital social, "
        "en qualité de {{ a.role }} ;\n"
        "{% endfor %}"
    )

    doc.add_paragraph()
    bp2 = doc.add_paragraph()
    bp2.add_run("IL A ÉTÉ CONVENU CE QUI SUIT :").bold = True
    doc.add_paragraph()

    _add_article(doc, 1, "OBJET ET DURÉE",
        "La Société a pour objet : {{ company_purpose }}.\n"
        "Elle est constituée pour une durée de {{ duration_years }} ans "
        "à compter de son immatriculation au {{ company_registry }}."
    )
    _add_article(doc, 2, "RÉPARTITION DU CAPITAL",
        "Le capital social est réparti comme suit :\n"
        "{% for a in associates %}"
        "- {{ a.name }} : {{ a.shares_percentage }}%\n"
        "{% endfor %}"
    )
    _add_article(doc, 3, "PRISE DE DÉCISION",
        "Les décisions ordinaires sont prises à la majorité simple (50%+1).\n"
        "Les décisions suivantes requièrent une majorité de {{ decisions_threshold }}% :\n"
        "- Modification des statuts ;\n"
        "- Augmentation ou réduction de capital ;\n"
        "- Fusion, scission ou dissolution ;\n"
        "- Cession de fonds de commerce."
    )
    _add_article(doc, 4, "DROIT DE PRÉEMPTION",
        "{% if right_of_first_refusal %}"
        "Tout associé souhaitant céder tout ou partie de ses parts doit en informer "
        "les autres associés par lettre recommandée avec accusé de réception. "
        "Les autres associés disposent d'un délai de 30 jours pour exercer leur droit "
        "de préemption, au prix proposé par le cessionnaire tiers."
        "{% else %}"
        "Les associés ne bénéficient pas d'un droit de préemption formel. "
        "Toute cession doit néanmoins être approuvée à l'unanimité des associés."
        "{% endif %}"
    )
    _add_article(doc, 5, "CLAUSE DE SORTIE CONJOINTE (TAG-ALONG)",
        "{% if tag_along %}"
        "En cas de cession de parts par un associé majoritaire à un tiers, "
        "les associés minoritaires ont le droit de céder leurs parts dans les mêmes "
        "conditions (prix, modalités) que l'associé cédant."
        "{% else %}"
        "Les parties renoncent expressément au droit de sortie conjointe pour la durée "
        "du présent pacte."
        "{% endif %}"
    )
    _add_article(doc, 6, "CLAUSE DE CESSION FORCÉE (DRAG-ALONG)",
        "{% if drag_along %}"
        "En cas d'offre de rachat par un tiers portant sur au moins 75% du capital, "
        "l'associé ou le groupe d'associés majoritaire peut exiger des associés minoritaires "
        "qu'ils cèdent leurs parts aux mêmes conditions."
        "{% else %}"
        "Les parties renoncent expressément à la clause de cession forcée pour la durée "
        "du présent pacte."
        "{% endif %}"
    )
    _add_article(doc, 7, "CONFIDENTIALITÉ",
        "Les associés s'engagent à maintenir la confidentialité du présent pacte "
        "et de toutes les informations relatives à la Société auxquelles ils auraient accès."
    )
    _add_article(doc, 8, "LOI APPLICABLE",
        "Le présent pacte est régi par le {{ ohada_reference }} "
        "et par le {{ governing_law_label }}. "
        "Tout litige sera soumis au {{ jurisdiction }}."
    )

    doc.add_paragraph(
        "Fait en autant d'exemplaires que de parties, "
        "à {{ signature_place|default('Abidjan') }}, le {{ signature_date }}."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Signatures des associés :").bold = True
    doc.add_paragraph(
        "{% for a in associates %}"
        "{{ a.name }} ({{ a.shares_percentage }}% – {{ a.role }}) : "
        "___________________________\n"
        "{% endfor %}"
    )
    _add_footer(doc)

    doc.save(os.path.join(TEMPLATES_DIR, "pacte_associes.docx"))
    print("  ✓ pacte_associes.docx")


# ── Main ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Génération des templates dans : {TEMPLATES_DIR}\n")
    create_cdi_template("CI")
    create_cdi_template("SN")
    create_cdd_template("CI")
    create_cdd_template("SN")
    create_nda_template()
    create_prestation_template()
    create_pacte_associes_template()
    print(f"\n✅ 7 templates générés avec succès dans {TEMPLATES_DIR}")
